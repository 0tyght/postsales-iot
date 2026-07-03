from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "docs" / "customer-local-installation-manual.md"
OUTPUT = ROOT / "docs" / "Post-Sales-IoT-Local-Server-Installation-Manual.docx"
FALLBACK_OUTPUT = ROOT / "docs" / "Post-Sales-IoT-Local-Server-Installation-Manual-v2.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text.strip())
    run.bold = bold
    run.font.name = "TH Sarabun New"
    run.font.size = Pt(13)


def add_table_from_rows(doc, rows):
    if not rows:
        return
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, bold=(row_idx == 0))
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx == 0:
                set_cell_shading(cells[idx], "E8EEF5")
    doc.add_paragraph()


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "TH Sarabun New"
    normal.font.size = Pt(14)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Heading 1", 22, "0B2545"),
        ("Heading 2", 18, "1F4D78"),
        ("Heading 3", 16, "1F4D78"),
    ]:
        style = doc.styles[name]
        style.font.name = "TH Sarabun New"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(90)
    run = p.add_run("คู่มือการติดตั้งและดูแลระบบ")
    run.font.name = "TH Sarabun New"
    run.font.size = Pt(30)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Post-Sales IoT Local Server Edition")
    run.font.name = "TH Sarabun New"
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("008C8C")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("สำหรับติดตั้งระบบบนเครื่องลูกค้า ใช้โดเมนลูกค้าเอง เชื่อม LINE OA และควบคุม License รายเดือน/รายปี")
    run.font.name = "TH Sarabun New"
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string("455A64")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    run = p.add_run("เวอร์ชันเอกสาร 1.0")
    run.font.name = "TH Sarabun New"
    run.font.size = Pt(14)

    doc.add_page_break()


def parse_markdown(doc, text):
    lines = text.splitlines()
    i = 0
    in_code = False
    code_lines = []
    table_rows = []

    def flush_code():
        nonlocal code_lines
        if code_lines:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            code_lines = []

    def flush_table():
        nonlocal table_rows
        if table_rows:
            cleaned = [r for r in table_rows if not all(re.match(r"^-+$", c.strip()) for c in r)]
            if len(cleaned) >= 1:
                add_table_from_rows(doc, cleaned)
            table_rows = []

    while i < len(lines):
        line = lines[i].rstrip()

        if line.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_table()
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if line.strip().startswith("|") and line.strip().endswith("|"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if not all(re.match(r"^:?-{3,}:?$", c) for c in cells):
                table_rows.append(cells)
            i += 1
            continue
        else:
            flush_table()

        if not line.strip():
            i += 1
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif re.match(r"^\d+\.\s+", line.strip()):
            p = doc.add_paragraph(style="List Number")
            p.add_run(re.sub(r"^\d+\.\s+", "", line.strip()))
        elif line.strip().startswith("- [ ]"):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run("☐ " + line.strip()[5:].strip())
        elif line.strip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line.strip()[2:])
        else:
            p = doc.add_paragraph()
            p.add_run(line)
        i += 1

    flush_code()
    flush_table()


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run("Post-Sales IoT Local Server Edition")
        run.font.name = "TH Sarabun New"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string("78909C")


def main():
    doc = Document()
    style_document(doc)
    add_cover(doc)
    parse_markdown(doc, SOURCE.read_text(encoding="utf-8"))
    add_footer(doc)
    try:
        doc.save(OUTPUT)
        print(OUTPUT)
    except PermissionError:
        doc.save(FALLBACK_OUTPUT)
        print(FALLBACK_OUTPUT)


if __name__ == "__main__":
    main()
